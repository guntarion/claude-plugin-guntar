#!/usr/bin/env python3
"""
Italicize English terms in a Word document based on CSV term list.

Usage:
    python 25-italicize-english-terms.py <terms.csv> <input.docx> [output.docx]
"""

import sys
import csv
import re
from pathlib import Path
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def read_terms_from_csv(csv_path):
    """
    Read terms from CSV file and sort by length (longest first).

    Args:
        csv_path: Path to CSV file with terms

    Returns:
        List of terms sorted by length (descending)
    """
    terms = []

    with open(csv_path, 'r', encoding='utf-8') as f:
        reader = csv.reader(f)
        for row in reader:
            if row and row[0].strip():
                term = row[0].strip()
                terms.append(term)

    # Sort by length (longest first) to avoid partial matches
    terms.sort(key=len, reverse=True)

    return terms


def is_heading_1(paragraph):
    """
    Check if paragraph is Heading 1 style.

    Args:
        paragraph: python-docx Paragraph object

    Returns:
        Boolean
    """
    if paragraph.style and paragraph.style.name:
        return 'Heading 1' in paragraph.style.name

    return False


def set_language_en_us(run):
    """
    Set language to English (US) for a run.

    Args:
        run: python-docx Run object
    """
    # Get or create rPr element
    rPr = run._element.get_or_add_rPr()

    # Get or create lang element
    lang = rPr.find(qn('w:lang'))

    if lang is None:
        lang = OxmlElement('w:lang')
        rPr.append(lang)

    # Set language to en-US
    lang.set(qn('w:val'), 'en-US')


def set_no_proofing(run):
    """
    Set 'Do not check spelling or grammar' for a run.

    Args:
        run: python-docx Run object
    """
    # Add noProof element
    rPr = run._element.get_or_add_rPr()
    noProof = rPr.find(qn('w:noProof'))

    if noProof is None:
        noProof = OxmlElement('w:noProof')
        rPr.append(noProof)


def process_paragraph(paragraph, terms):
    """
    Find English terms in paragraph and italicize them.

    Args:
        paragraph: python-docx Paragraph object
        terms: List of English terms sorted by length

    Returns:
        Number of terms found and italicized
    """
    # Skip Heading 1
    if is_heading_1(paragraph):
        return 0

    # Get full paragraph text
    full_text = paragraph.text

    if not full_text.strip():
        return 0

    # Find all term occurrences (case-insensitive)
    term_positions = []  # List of (start, end, term)

    for term in terms:
        # Use word boundary to match whole words/phrases
        # Escape special regex characters
        escaped_term = re.escape(term)
        pattern = r'\b' + escaped_term + r'\b'

        for match in re.finditer(pattern, full_text, re.IGNORECASE):
            start = match.start()
            end = match.end()
            matched_text = full_text[start:end]
            term_positions.append((start, end, matched_text))

    if not term_positions:
        return 0

    # Sort by start position
    term_positions.sort(key=lambda x: x[0])

    # Remove overlapping matches (keep first match)
    filtered_positions = []
    last_end = -1

    for start, end, text in term_positions:
        if start >= last_end:
            filtered_positions.append((start, end, text))
            last_end = end

    if not filtered_positions:
        return 0

    # Build character-to-formatting map from existing runs
    char_to_format = []

    for run in paragraph.runs:
        run_len = len(run.text)
        for i in range(run_len):
            char_to_format.append({
                'bold': run.bold,
                'italic': run.italic,
                'underline': run.underline,
                'font_name': run.font.name if run.font.name else None,
                'font_size': run.font.size,
                'font_color': run.font.color.rgb if run.font.color.rgb else None,
            })

    # Clear all runs
    for run in paragraph.runs:
        run.text = ""

    # Rebuild paragraph with terms italicized
    current_pos = 0
    current_run = None
    term_run = None

    for start, end, matched_text in filtered_positions:
        # Add text before term (normal formatting)
        if start > current_pos:
            before_text = full_text[current_pos:start]

            for i, char in enumerate(before_text):
                char_idx = current_pos + i
                if char_idx < len(char_to_format):
                    char_format = char_to_format[char_idx]
                else:
                    char_format = {}

                # Check if we need new run (format changed)
                if current_run is None or not runs_have_same_format(current_run, char_format):
                    # Start new run
                    current_run = paragraph.add_run(char)
                    apply_format_to_run(current_run, char_format)
                else:
                    # Add to current run
                    current_run.text += char

        # Add term (with italic + English + no proofing)
        term_run = paragraph.add_run(matched_text)

        # Get base format from first character of term
        if start < len(char_to_format):
            base_format = char_to_format[start]
            apply_format_to_run(term_run, base_format)

        # Apply English term formatting
        term_run.italic = True
        set_language_en_us(term_run)
        set_no_proofing(term_run)

        current_pos = end
        current_run = None  # Force new run after term

    # Add remaining text after last term
    if current_pos < len(full_text):
        remaining_text = full_text[current_pos:]

        for i, char in enumerate(remaining_text):
            char_idx = current_pos + i
            if char_idx < len(char_to_format):
                char_format = char_to_format[char_idx]
            else:
                char_format = {}

            if current_run is None or not runs_have_same_format(current_run, char_format):
                current_run = paragraph.add_run(char)
                apply_format_to_run(current_run, char_format)
            else:
                current_run.text += char

    return len(filtered_positions)


def runs_have_same_format(run, char_format):
    """Check if run has same format as char_format dict."""
    return (
        run.bold == char_format.get('bold') and
        run.italic == char_format.get('italic') and
        run.underline == char_format.get('underline') and
        run.font.name == char_format.get('font_name') and
        run.font.size == char_format.get('font_size')
    )


def apply_format_to_run(run, char_format):
    """Apply formatting from char_format dict to run."""
    run.bold = char_format.get('bold')
    run.italic = char_format.get('italic')
    run.underline = char_format.get('underline')

    if char_format.get('font_name'):
        run.font.name = char_format['font_name']
    if char_format.get('font_size'):
        run.font.size = char_format['font_size']
    if char_format.get('font_color'):
        run.font.color.rgb = char_format['font_color']


def process_document(doc, terms):
    """
    Process entire document - paragraphs and tables.

    Args:
        doc: python-docx Document object
        terms: List of English terms

    Returns:
        Total count of terms italicized
    """
    total_terms = 0

    # Process paragraphs
    for paragraph in doc.paragraphs:
        count = process_paragraph(paragraph, terms)
        total_terms += count

    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    count = process_paragraph(paragraph, terms)
                    total_terms += count

    return total_terms


def main():
    if len(sys.argv) < 3:
        print("Usage:")
        print("  python 25-italicize-english-terms.py <terms.csv> <input.docx> [output.docx]")
        print("\nExamples:")
        print("  python 25-italicize-english-terms.py terms.csv document.docx")
        print("  python 25-italicize-english-terms.py terms.csv input.docx output.docx")
        sys.exit(1)

    csv_file = Path(sys.argv[1])
    input_file = Path(sys.argv[2])

    # Validate files
    if not csv_file.exists():
        print(f"Error: CSV file not found: {csv_file}")
        sys.exit(1)

    if not input_file.exists():
        print(f"Error: Input file not found: {input_file}")
        sys.exit(1)

    if input_file.suffix.lower() != '.docx':
        print(f"Error: Input must be a .docx file: {input_file}")
        sys.exit(1)

    # Determine output file
    if len(sys.argv) >= 4:
        output_file = Path(sys.argv[3])
        if not output_file.suffix:
            output_file = output_file.with_suffix('.docx')
    else:
        # Generate output filename: input-italicized.docx
        output_file = input_file.parent / f"{input_file.stem}-italicized.docx"

    print(f"Terms CSV: {csv_file}")
    print(f"Input:     {input_file}")
    print(f"Output:    {output_file}")
    print()

    # Read terms
    print("Reading English terms from CSV...")
    terms = read_terms_from_csv(csv_file)
    print(f"Found {len(terms)} terms (sorted by length):")
    for i, term in enumerate(terms[:10], 1):
        print(f"  {i}. '{term}' ({len(term)} chars)")
    if len(terms) > 10:
        print(f"  ... and {len(terms) - 10} more")
    print()

    # Load document
    print("Loading document...")
    doc = Document(str(input_file))

    total_paragraphs = len(doc.paragraphs)
    total_tables = len(doc.tables)
    print(f"Found {total_paragraphs} paragraphs and {total_tables} tables")

    # Process document
    print("\nItalicizing English terms...")
    total_terms = process_document(doc, terms)

    # Save output
    doc.save(str(output_file))

    print(f"\n{'='*60}")
    print(f"✓ Processing complete!")
    print(f"  Terms italicized: {total_terms} occurrences")
    print(f"  Output saved: {output_file}")
    print(f"{'='*60}")


if __name__ == "__main__":
    main()
