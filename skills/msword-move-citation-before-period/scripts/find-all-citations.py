#!/usr/bin/env python3
"""Find ALL citations in entire document."""

import sys
import re
from pathlib import Path
from docx import Document

def main():
    if len(sys.argv) < 2:
        print("Usage: python find-all-citations.py <file.docx>")
        sys.exit(1)

    doc_path = Path(sys.argv[1])
    doc = Document(str(doc_path))

    print(f"Analyzing: {doc_path.name}")
    print(f"Total paragraphs: {len(doc.paragraphs)}\n")

    # Pattern for any bracketed number
    pattern = r'\[\d{1,3}\]'

    found_citations = []

    # Search all paragraphs
    for idx, para in enumerate(doc.paragraphs):
        text = para.text
        matches = re.findall(pattern, text)

        if matches:
            found_citations.extend(matches)
            print(f"Para {idx}: {len(matches)} citations")
            # Show context around citation
            for match in matches:
                pos = text.find(match)
                start = max(0, pos - 30)
                end = min(len(text), pos + len(match) + 30)
                context = text[start:end]
                print(f"  Context: ...{context}...")
            print()

    # Search tables
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                for para in cell.paragraphs:
                    text = para.text
                    matches = re.findall(pattern, text)
                    if matches:
                        found_citations.extend(matches)
                        print(f"Table {table_idx}, Row {row_idx}, Cell {cell_idx}: {len(matches)} citations")
                        for match in matches:
                            pos = text.find(match)
                            start = max(0, pos - 30)
                            end = min(len(text), pos + len(match) + 30)
                            context = text[start:end]
                            print(f"  Context: ...{context}...")

    print(f"\n{'='*60}")
    print(f"TOTAL CITATIONS FOUND: {len(found_citations)}")
    print(f"Citations: {found_citations}")
    print(f"{'='*60}")

if __name__ == "__main__":
    main()
