#!/usr/bin/env python3
"""
Move citation numbers from after period to before period with space.

Transforms: "text.[123]" → "text [123]."

Usage:
    python 24-move-citations-before-period.py <input.docx> [output.docx]
"""

import sys
import re
from pathlib import Path
from docx import Document


def move_citations_in_text(text):
    """
    Move citations from after period to before period with space.

    Pattern: "text.[123]" → "text [123]."
    Handles 1-3 digit citation numbers.

    Args:
        text: String to process

    Returns:
        Processed string with citations moved
    """
    # Pattern: period followed by bracketed number (1-3 digits)
    # Example: "Indonesia.[14]" → "Indonesia [14]."
    pattern = r'\.(\[\d{1,3}\])'
    replacement = r' \1.'

    return re.sub(pattern, replacement, text)


def process_paragraph(paragraph):
    """
    Process a paragraph, moving citations while preserving formatting.

    Args:
        paragraph: python-docx Paragraph object

    Returns:
        True if paragraph was modified, False otherwise
    """
    # Get full paragraph text
    original_text = paragraph.text

    # Check if any citations need moving
    if not re.search(r'\.(\[\d{1,3}\])', original_text):
        return False

    # Apply transformation
    new_text = move_citations_in_text(original_text)

    if original_text == new_text:
        return False

    # Build mapping of character position to formatting (from runs)
    char_to_format = []
    char_pos = 0

    for run in paragraph.runs:
        run_len = len(run.text)
        for i in range(run_len):
            char_to_format.append({
                'bold': run.bold,
                'italic': run.italic,
                'underline': run.underline,
                'font_name': run.font.name if run.font.name else None,
                'font_size': run.font.size,
                'font_color': run.font.color.rgb if run.font.color.rgb else None,
            })
        char_pos += run_len

    # Map new text characters to old text positions
    # The transformation is ".[n]" (4-6 chars) → " [n]." (5-7 chars)
    # So we need to track which new chars correspond to which old chars

    # Find all pattern matches in original text
    pattern = r'\.(\[\d{1,3}\])'
    old_to_new_map = {}  # Maps old char index to new char index

    offset = 0  # Tracks cumulative length change
    for match in re.finditer(pattern, original_text):
        match_start = match.start()
        match_end = match.end()

        # Before the match, mapping is identity + offset
        for i in range(len(old_to_new_map), match_start):
            old_to_new_map[i] = i + offset

        # The match: ".[123]" → " [123]."
        # Period moves from position match_start to match_end-1+offset
        # Bracket part stays roughly the same position
        old_to_new_map[match_start] = match_end - 1 + offset  # Period goes to end
        for i in range(match_start + 1, match_end):
            old_to_new_map[i] = match_start + offset + (i - match_start)  # Brackets shift left by 1

        offset += 1  # Each replacement adds 1 character (space)

    # Map remaining characters
    for i in range(len(old_to_new_map), len(original_text)):
        old_to_new_map[i] = i + offset

    # Create reverse mapping: new char index → old char index
    new_to_old_map = {}
    for old_idx, new_idx in old_to_new_map.items():
        new_to_old_map[new_idx] = old_idx

    # Fill in gaps in new_to_old_map (for the added spaces)
    for i in range(len(new_text)):
        if i not in new_to_old_map:
            # This is a new character (space), use format from preceding char
            if i > 0 and (i-1) in new_to_old_map:
                new_to_old_map[i] = new_to_old_map[i-1]
            elif i+1 < len(new_text) and (i+1) in new_to_old_map:
                new_to_old_map[i] = new_to_old_map[i+1]

    # Clear all runs
    for run in paragraph.runs:
        run.text = ""

    # Rebuild runs with proper formatting
    current_format = None
    current_run = None
    current_text = ""

    for i, char in enumerate(new_text):
        # Get format for this character
        old_idx = new_to_old_map.get(i, min(len(char_to_format)-1, i))
        old_idx = min(old_idx, len(char_to_format)-1)

        if old_idx < len(char_to_format):
            char_format = char_to_format[old_idx]
        else:
            char_format = char_to_format[-1] if char_to_format else {}

        # Check if format changed
        if current_format is None or current_format != char_format:
            # Save previous run
            if current_run is not None and current_text:
                current_run.text = current_text

            # Start new run
            current_run = paragraph.add_run()
            current_run.bold = char_format.get('bold')
            current_run.italic = char_format.get('italic')
            current_run.underline = char_format.get('underline')
            if char_format.get('font_name'):
                current_run.font.name = char_format['font_name']
            if char_format.get('font_size'):
                current_run.font.size = char_format['font_size']
            if char_format.get('font_color'):
                current_run.font.color.rgb = char_format['font_color']

            current_format = char_format
            current_text = char

        else:
            # Same format, add to current run text
            current_text += char

    # Save final run
    if current_run is not None and current_text:
        current_run.text = current_text

    return True


def process_document(doc):
    """
    Process entire document - paragraphs and table cells.

    Args:
        doc: python-docx Document object

    Returns:
        Count of modifications made
    """
    modifications = 0

    # Process paragraphs
    for paragraph in doc.paragraphs:
        if process_paragraph(paragraph):
            modifications += 1

    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if process_paragraph(paragraph):
                        modifications += 1

    return modifications


def main():
    if len(sys.argv) < 2:
        print("Usage:")
        print("  python 24-move-citations-before-period.py <input.docx> [output.docx]")
        print("\nExamples:")
        print("  python 24-move-citations-before-period.py document.docx")
        print("  python 24-move-citations-before-period.py input.docx output.docx")
        print("\nTransforms:")
        print('  "text.[123]" → "text [123]."')
        print('  "word.[5]" → "word [5]."')
        sys.exit(1)

    input_file = Path(sys.argv[1])

    # Validate input file
    if not input_file.exists():
        print(f"Error: File not found: {input_file}")
        sys.exit(1)

    if input_file.suffix.lower() != '.docx':
        print(f"Error: File must be a .docx file: {input_file}")
        sys.exit(1)

    # Determine output file
    if len(sys.argv) >= 3:
        output_file = Path(sys.argv[2])
        if not output_file.suffix:
            output_file = output_file.with_suffix('.docx')
    else:
        # Generate output filename: input-fixed.docx
        output_file = input_file.parent / f"{input_file.stem}-fixed.docx"

    print(f"Input:  {input_file}")
    print(f"Output: {output_file}")
    print()

    # Load document
    print("Loading document...")
    doc = Document(str(input_file))

    # Count total elements
    total_paragraphs = len(doc.paragraphs)
    total_tables = len(doc.tables)
    print(f"Found {total_paragraphs} paragraphs and {total_tables} tables")

    # Process document
    print("\nProcessing citations...")
    modifications = process_document(doc)

    # Save output
    doc.save(str(output_file))

    print(f"\n{'='*60}")
    print(f"✓ Processing complete!")
    print(f"  Modified elements: {modifications}")
    print(f"  Output saved: {output_file}")
    print(f"{'='*60}")


if __name__ == "__main__":
    main()
