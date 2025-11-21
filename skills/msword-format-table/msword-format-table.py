#!/usr/bin/env python3
"""
Format tables in Word documents:
1. Center align headers (horizontal and vertical) + bold
2. Remove bold from first column body (excluding header)
3. Left align all body content (excluding headers)
4. Set repeat header rows
5. Apply table style

Usage:
    python 26-format-tables.py <input.docx> [output.docx] [--style "Table Grid"]
"""

import sys
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def center_align_headers(table):
    """
    Center align header row (first row) both horizontally and vertically, and make bold.

    Args:
        table: python-docx Table object

    Returns:
        Number of header cells formatted
    """
    if not table.rows:
        return 0

    header_count = 0
    first_row = table.rows[0]

    for cell in first_row.cells:
        # Horizontal alignment
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Make bold
            for run in paragraph.runs:
                run.bold = True

        # Vertical alignment
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        header_count += 1

    return header_count


def set_repeat_header_rows(table):
    """
    Set header row to repeat on each page.

    Args:
        table: python-docx Table object

    Returns:
        True if set successfully
    """
    if not table.rows:
        return False

    try:
        # Set first row to repeat as header
        first_row = table.rows[0]
        tr = first_row._tr
        trPr = tr.get_or_add_trPr()

        # Add tblHeader element if it doesn't exist
        tblHeader = trPr.find(qn('w:tblHeader'))
        if tblHeader is None:
            tblHeader = OxmlElement('w:tblHeader')
            trPr.append(tblHeader)

        return True
    except:
        return False


def set_table_style(table, style_name):
    """
    Apply a table style.

    Args:
        table: python-docx Table object
        style_name: Name of table style (e.g., "Table Grid", "Grid Table 4 - Accent 4")

    Returns:
        True if style applied successfully
    """
    if not style_name:
        return False

    try:
        # Try the style name as-is
        table.style = style_name
        return True
    except:
        # Try without spaces (some styles use camelCase)
        try:
            table.style = style_name.replace(' ', '')
            return True
        except:
            # Try with "Table" prefix if not present
            try:
                if not style_name.startswith('Table'):
                    table.style = 'Table' + style_name
                    return True
            except:
                pass

    return False


def unbold_first_column_body(table):
    """
    Remove bold from first column (excluding header row).

    Args:
        table: python-docx Table object

    Returns:
        Number of cells unbolded
    """
    if len(table.rows) <= 1:
        return 0  # No body rows

    unbold_count = 0

    # Process rows starting from second row (skip header)
    for row_idx, row in enumerate(table.rows[1:], start=1):
        if row.cells:
            # Get first cell
            first_cell = row.cells[0]

            # Remove bold from all runs in this cell
            for paragraph in first_cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = False

            unbold_count += 1

    return unbold_count


def left_align_body(table):
    """
    Left align all body content (excluding header row).

    Args:
        table: python-docx Table object

    Returns:
        Number of cells aligned
    """
    if len(table.rows) <= 1:
        return 0  # No body rows

    align_count = 0

    # Process rows starting from second row (skip header)
    for row in table.rows[1:]:
        for cell in row.cells:
            # Left align all paragraphs in cell
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

            align_count += 1

    return align_count


def format_all_tables(doc, table_style='Table Grid'):
    """
    Apply all formatting to all tables in document.

    Args:
        doc: python-docx Document object
        table_style: Table style name to apply (default: 'Table Grid')

    Returns:
        Dictionary with counts
    """
    stats = {
        'tables_found': 0,
        'headers_centered': 0,
        'headers_bolded': 0,
        'repeat_header_set': 0,
        'style_applied': 0,
        'first_col_unbolded': 0,
        'body_cells_aligned': 0
    }

    for table in doc.tables:
        stats['tables_found'] += 1

        # 1. Apply table style
        if set_table_style(table, table_style):
            stats['style_applied'] += 1

        # 2. Center align headers and make bold
        header_count = center_align_headers(table)
        stats['headers_centered'] += header_count
        if header_count > 0:
            stats['headers_bolded'] += 1

        # 3. Set repeat header rows
        if set_repeat_header_rows(table):
            stats['repeat_header_set'] += 1

        # 4. Unbold first column body
        unbold_count = unbold_first_column_body(table)
        stats['first_col_unbolded'] += unbold_count

        # 5. Left align body
        align_count = left_align_body(table)
        stats['body_cells_aligned'] += align_count

    return stats


def main():
    if len(sys.argv) < 2:
        print("Usage:")
        print("  python 26-format-tables.py <input.docx> [output.docx] [--style \"Table Grid\"]")
        print("\nExamples:")
        print("  python 26-format-tables.py document.docx")
        print("  python 26-format-tables.py input.docx output.docx")
        print("  python 26-format-tables.py input.docx --style \"Grid Table 4 - Accent 4\"")
        print("  python 26-format-tables.py input.docx output.docx --style \"Table Grid\"")
        print("\nOperations:")
        print("  1. Apply table style (default: 'Table Grid')")
        print("  2. Center align table headers (horizontal + vertical)")
        print("  3. Make header text bold")
        print("  4. Set repeat header rows")
        print("  5. Remove bold from first column body")
        print("  6. Left align all body content")
        sys.exit(1)

    # Parse arguments
    args = sys.argv[1:]
    input_file = None
    output_file = None
    table_style = 'Table Grid'  # Default style

    # Parse --style argument
    if '--style' in args:
        style_idx = args.index('--style')
        if style_idx + 1 < len(args):
            table_style = args[style_idx + 1]
            # Remove --style and its value from args
            args.pop(style_idx)  # Remove --style
            args.pop(style_idx)  # Remove style value

    # Remaining args are input and output files
    if len(args) >= 1:
        input_file = Path(args[0])

    if len(args) >= 2:
        output_file = Path(args[1])

    if not input_file:
        print("Error: Input file required")
        sys.exit(1)

    # Validate input file
    if not input_file.exists():
        print(f"Error: File not found: {input_file}")
        sys.exit(1)

    if input_file.suffix.lower() != '.docx':
        print(f"Error: File must be a .docx file: {input_file}")
        sys.exit(1)

    # Determine output file
    if not output_file:
        # Generate output filename: input-formatted.docx
        output_file = input_file.parent / f"{input_file.stem}-formatted.docx"
    elif not output_file.suffix:
        output_file = output_file.with_suffix('.docx')

    print(f"Input:  {input_file}")
    print(f"Output: {output_file}")
    print(f"Style:  {table_style}")
    print()

    # Load document
    print("Loading document...")
    doc = Document(str(input_file))

    # Format tables
    print("Formatting tables...")
    stats = format_all_tables(doc, table_style)

    # Save output
    doc.save(str(output_file))

    print(f"\n{'='*60}")
    print(f"✓ Processing complete!")
    print(f"  Tables found: {stats['tables_found']}")
    print(f"  Style applied: {stats['style_applied']} tables")
    print(f"  Headers centered: {stats['headers_centered']} cells")
    print(f"  Headers bolded: {stats['headers_bolded']} tables")
    print(f"  Repeat header set: {stats['repeat_header_set']} tables")
    print(f"  First column unbolded: {stats['first_col_unbolded']} cells")
    print(f"  Body cells left-aligned: {stats['body_cells_aligned']} cells")
    print(f"  Output saved: {output_file}")
    print(f"{'='*60}")


if __name__ == "__main__":
    main()
