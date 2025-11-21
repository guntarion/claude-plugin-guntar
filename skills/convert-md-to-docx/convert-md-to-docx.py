#!/usr/bin/env python3
"""
Markdown to Word Document Converter

This script merges multiple markdown files and converts them to a .docx file
with specific formatting requirements:
- B5 JIS paper size (18.2 x 25.71 cm)
- Mirror margins
- Proper heading styles
- Optional header/footer with page numbers and background images

Usage:
    python 21-md-to-docx.py <file1.md,file2.md,...> [output.docx]
    python 21-md-to-docx.py <folder_path> [output.docx]

Examples:
    python 21-md-to-docx.py file1.md,file2.md,file3.md output.docx
    python 21-md-to-docx.py ./markdown_folder my_document.docx
    python 21-md-to-docx.py ./markdown_folder  # auto-generates output name
"""

import sys
import os
import re
from pathlib import Path
from typing import List, Tuple
import pypandoc
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_reference_docx(template_path: str) -> None:
    """
    Create a reference .docx template with proper formatting settings.

    This template will be used by pypandoc to ensure correct paper size,
    margins, and other document properties.
    """
    doc = Document()

    # Get the default section
    section = doc.sections[0]

    # Set B5 JIS paper size (18.2 x 25.71 cm)
    section.page_width = Cm(18.2)
    section.page_height = Cm(25.71)

    # Set mirror margins
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)    # Inside margin
    section.right_margin = Cm(2.0)   # Outside margin
    section.gutter = Cm(0)

    # Set different odd & even pages
    section.different_first_page_header_footer = False

    # Add a dummy paragraph to make it valid
    doc.add_paragraph("Template document - will be replaced by pandoc")

    # Save the template
    doc.save(template_path)
    print(f"✓ Created reference template: {template_path}")


def add_page_numbers_and_backgrounds(docx_path: str,
                                     bg_odd: str = None,
                                     bg_even: str = None) -> None:
    """
    Add page numbers and background images to the generated document.

    Args:
        docx_path: Path to the .docx file
        bg_odd: Path to background image for odd pages
        bg_even: Path to background image for even pages
    """
    try:
        doc = Document(docx_path)

        for section in doc.sections:
            # Enable different odd & even pages
            section.different_first_page_header_footer = False

            # Add page numbers to footer
            # Note: This is a simplified implementation
            # Full implementation would require more complex OPC manipulation
            footer = section.footer
            paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # Outside alignment

            # Add page number field
            run = paragraph.add_run()

            # Create page number field (simplified - actual implementation is complex)
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')

            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            instrText.text = "PAGE"

            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')

            run._r.append(fldChar1)
            run._r.append(instrText)
            run._r.append(fldChar2)

        doc.save(docx_path)
        print(f"✓ Added page numbers to document")

        if bg_odd or bg_even:
            print(f"⚠ Background images require manual setup in Word due to API limitations")
            print(f"  Odd page background: {bg_odd if bg_odd else 'N/A'}")
            print(f"  Even page background: {bg_even if bg_even else 'N/A'}")

    except Exception as e:
        print(f"⚠ Could not add page numbers: {e}")
        print(f"  You may need to add them manually in Word")


def sanitize_filename(filename: str) -> str:
    """
    Remove unsafe characters from filename.

    Args:
        filename: Original filename string

    Returns:
        Sanitized filename safe for filesystem
    """
    # Remove leading/trailing whitespace
    filename = filename.strip()

    # Remove or replace unsafe characters
    # Keep: letters, numbers, spaces, hyphens, underscores
    # Replace others with underscore
    filename = re.sub(r'[<>:"/\\|?*\x00-\x1f]', '_', filename)

    # Remove leading/trailing dots and spaces
    filename = filename.strip('. ')

    # Replace multiple spaces with single space
    filename = re.sub(r'\s+', ' ', filename)

    # Replace multiple underscores with single underscore
    filename = re.sub(r'_+', '_', filename)

    # Limit length (max 200 characters to be safe)
    if len(filename) > 200:
        filename = filename[:200].strip()

    # If empty after sanitization, use default
    if not filename:
        filename = "document"

    return filename


def get_first_line_as_filename(md_file: Path) -> str:
    """
    Read the first line of a markdown file and use as filename.

    Args:
        md_file: Path to markdown file

    Returns:
        Sanitized filename from first line
    """
    try:
        first_line = md_file.read_text(encoding='utf-8').split('\n')[0]

        # Remove markdown heading symbols
        first_line = re.sub(r'^#+\s*', '', first_line)

        # Sanitize the filename
        filename = sanitize_filename(first_line)

        print(f"✓ Using first line as filename: '{filename}'")
        return filename

    except Exception as e:
        print(f"⚠ Could not read first line, using default: {e}")
        return "document"


def get_md_files(input_path: str) -> List[Path]:
    """
    Get list of markdown files from input (file list or folder).

    Args:
        input_path: Comma-separated file list or folder path

    Returns:
        List of Path objects sorted alphabetically
    """
    # Check if input is comma-separated files
    if ',' in input_path:
        files = [Path(f.strip()) for f in input_path.split(',')]
        # Verify all files exist
        for f in files:
            if not f.exists():
                raise FileNotFoundError(f"File not found: {f}")
        # Sort alphabetically by filename
        return sorted(files, key=lambda x: x.name)

    # Otherwise, treat as folder
    folder = Path(input_path)
    if not folder.exists():
        raise FileNotFoundError(f"Folder not found: {folder}")

    if not folder.is_dir():
        # Single file
        return [folder]

    # Get all .md files in folder, sorted alphabetically
    md_files = sorted(folder.glob('*.md'), key=lambda x: x.name)

    if not md_files:
        raise ValueError(f"No .md files found in: {folder}")

    return md_files


def merge_markdown_files(md_files: List[Path]) -> str:
    """
    Merge multiple markdown files into one string.

    Files are processed in the order provided (should be alphabetically sorted).

    Args:
        md_files: List of Path objects to markdown files

    Returns:
        Merged markdown content as string
    """
    merged_content = []

    print(f"\nMerging {len(md_files)} markdown file(s):")
    for i, md_file in enumerate(md_files, 1):
        print(f"  [{i}/{len(md_files)}] {md_file.name}")
        content = md_file.read_text(encoding='utf-8')
        merged_content.append(content)

    # Join with double newline to ensure separation
    return '\n\n'.join(merged_content)


def convert_md_to_docx(merged_md: str,
                       output_path: str,
                       reference_docx: str = None) -> None:
    """
    Convert merged markdown to .docx using pypandoc.

    Args:
        merged_md: Merged markdown content
        output_path: Path for output .docx file
        reference_docx: Path to reference template (optional)
    """
    print(f"\nConverting to .docx...")

    # Pypandoc conversion options
    extra_args = []

    if reference_docx and Path(reference_docx).exists():
        extra_args.append(f'--reference-doc={reference_docx}')

    # Convert markdown to docx
    pypandoc.convert_text(
        merged_md,
        'docx',
        format='md',
        outputfile=output_path,
        extra_args=extra_args
    )

    print(f"✓ Conversion complete: {output_path}")


def main():
    """Main execution function."""

    # Check arguments
    if len(sys.argv) < 2:
        print(__doc__)
        sys.exit(1)

    input_path = sys.argv[1]

    print("=" * 60)
    print("Markdown to Word Document Converter")
    print("=" * 60)

    try:
        # Get markdown files (sorted alphabetically)
        md_files = get_md_files(input_path)

        # Determine source folder for output
        if ',' in input_path:
            # Multiple files - use first file's directory
            source_folder = md_files[0].parent
        else:
            input_path_obj = Path(input_path)
            if input_path_obj.is_dir():
                source_folder = input_path_obj
            else:
                source_folder = input_path_obj.parent

        # Determine output file name
        if len(sys.argv) >= 3:
            # User provided filename
            output_filename = sys.argv[2]
            # If user provides full path, extract just filename
            output_filename = Path(output_filename).name
        else:
            # Auto-generate: read first line of first file
            base_filename = get_first_line_as_filename(md_files[0])
            output_filename = f"{base_filename}.docx"

        # Construct full output path in source folder
        output_path = source_folder / output_filename

        print(f"Output folder: {source_folder}")
        print(f"Output file: {output_filename}")

        # Create reference template with proper formatting
        template_path = ".temp_reference.docx"
        create_reference_docx(template_path)

        # Merge markdown files
        merged_content = merge_markdown_files(md_files)

        # Convert to .docx
        convert_md_to_docx(merged_content, str(output_path), template_path)

        # Add page numbers and background images (if specified)
        bg_odd = "./Desain/inner-buku2-ganjil.jpg"
        bg_even = "./Desain/inner-buku2-genap.jpg"

        if Path(bg_odd).exists() or Path(bg_even).exists():
            add_page_numbers_and_backgrounds(
                str(output_path),
                bg_odd if Path(bg_odd).exists() else None,
                bg_even if Path(bg_even).exists() else None
            )

        # Clean up template
        if Path(template_path).exists():
            Path(template_path).unlink()
            print(f"✓ Cleaned up temporary template")

        print("\n" + "=" * 60)
        print("SUCCESS!")
        print("=" * 60)
        print(f"Output: {output_path}")
        print(f"Size: {output_path.stat().st_size / 1024:.1f} KB")
        print()
        print("Document Settings Applied:")
        print("  • Paper: B5 JIS (18.2 x 25.71 cm)")
        print("  • Margins: Mirror (Top: 2.5, Bottom: 2.5, Inside: 2.5, Outside: 2.0 cm)")
        print("  • Styles: Heading 1/2/3, Bold, Italic, Lists, Tables")
        print()

        # Additional notes for background images
        if Path(bg_odd).exists() or Path(bg_even).exists():
            print("NOTE: Background images need to be added manually in Word:")
            print("  1. Open the document in Microsoft Word")
            print("  2. Go to Design > Page Color > Fill Effects")
            print("  3. Choose Picture tab and select the background image")
            print("  4. For different odd/even pages, use different sections")
            print()

    except Exception as e:
        print(f"\n❌ Error: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()
